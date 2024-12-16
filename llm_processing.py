import os
import whisper
from dotenv import load_dotenv
from llama_index.core import SummaryIndex, Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_index.llms.groq import Groq
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from docx import Document
from llama_index.core import SimpleDirectoryReader  # Assuming you need this for reading files
import nest_asyncio
import warnings

# Suppress all warnings
warnings.filterwarnings("ignore")

# Initialize nest_asyncio
nest_asyncio.apply()

# Set up the environment and Groq model for summarization
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
os.environ["GROQ_API_KEY"] = GROQ_API_KEY

llm = Groq(model="llama-3.1-8b-instant")
Settings.llm = llm
Settings.embed_model = HuggingFaceEmbedding()

# Function to transcribe audio
def transcribe_audio(audio_file_path):
    model = whisper.load_model("base")
    result = model.transcribe(audio_file_path)
    transcribed_text = result['text']
    return transcribed_text

# Function to create .docx from transcribed text
def create_docx_from_text(text, file_name="transcription.docx"):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_name)
    print(f"Transcription saved as '{file_name}'.")

# Main process for transcribing, saving, and summarizing
def main():
    audio_file_path = input("Enter the audio file path: ")
    print("Transcribing audio...")
    transcribed_text = transcribe_audio(audio_file_path)
    print("Transcription complete!")

    # Save the transcription to a .docx file
    docx_file = "transcription.docx"
    print("Creating document from transcription...")
    create_docx_from_text(transcribed_text, docx_file)


    documents = SimpleDirectoryReader(input_files=[docx_file]).load_data()


    splitter = SentenceSplitter(chunk_size=2048)
    nodes = splitter.get_nodes_from_documents(documents)


    summary_index = SummaryIndex(nodes)

    # Step 8: Configure the query engine
    summary_query_engine = summary_index.as_query_engine(
        response_mode="tree_summarize",
        use_async=True
    )


    response = summary_query_engine.query("Give a detailed summary of the given document.")

    summary_text = response.content if hasattr(response, 'content') else str(response)


    summary_index.storage_context.persist("summary_index")

    doc = Document()
    doc.add_heading("Minutes Of Meeting", level=1)
    doc.add_paragraph(summary_text)  # Add the generated summary as content.
    doc.save("outputs/MoM.docx")
    print("Summary saved as 'MoM.docx'.")

if __name__ == "__main__":
    main()


